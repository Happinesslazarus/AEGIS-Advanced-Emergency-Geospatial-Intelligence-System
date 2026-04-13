"""Generate AEGIS-Technical-With-Explanations as DOCX and PDF."""
import pathlib, textwrap
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown, subprocess, os

SRC = pathlib.Path(__file__).with_name("AEGIS-Technical-With-Explanations.md")
STEM = SRC.stem

# ── Word ──────────────────────────────────────────────────────────────
doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)

GREEN = RGBColor(0x1B, 0x7A, 0x3D)
BLUE  = RGBColor(0x1A, 0x56, 0xDB)
GRAY  = RGBColor(0x55, 0x55, 0x55)

def add_heading(text, level):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = GREEN if level <= 2 else BLUE

lines = SRC.read_text(encoding="utf-8").splitlines()
i = 0
while i < len(lines):
    line = lines[i]

    # Headings
    if line.startswith("# "):
        add_heading(line[2:].strip(), 1)
        i += 1; continue
    if line.startswith("## "):
        add_heading(line[3:].strip(), 2)
        i += 1; continue
    if line.startswith("### "):
        add_heading(line[4:].strip(), 3)
        i += 1; continue

    # Horizontal rule → skip
    if line.strip() == "---":
        i += 1; continue

    # Table
    if line.startswith("|"):
        # collect table lines
        tbl_lines = []
        while i < len(lines) and lines[i].startswith("|"):
            tbl_lines.append(lines[i]); i += 1
        # parse
        rows = []
        for tl in tbl_lines:
            cells = [c.strip() for c in tl.strip("|").split("|")]
            if all(set(c) <= set("- :") for c in cells):
                continue  # separator line
            rows.append(cells)
        if rows:
            ncols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=ncols, style="Light Grid Accent 1")
            for ri, row in enumerate(rows):
                for ci, cell in enumerate(row):
                    table.rows[ri].cells[ci].text = cell
                    for p in table.rows[ri].cells[ci].paragraphs:
                        p.style.font.size = Pt(9)
            if ri == 0:
                for ci in range(ncols):
                    for p in table.rows[0].cells[ci].paragraphs:
                        for run in p.runs:
                            run.bold = True
        continue

    # Block quote (> text) → indented, italic, colored
    if line.startswith("> "):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.4)
        run = p.add_run(line[2:].strip().replace("**", ""))
        run.italic = True
        run.font.color.rgb = GRAY
        i += 1; continue

    # Bullet points
    if line.startswith("- **") or line.startswith("   - ") or line.startswith("  - "):
        p = doc.add_paragraph(style="List Bullet")
        # Bold part
        stripped = line.lstrip(" -").strip()
        if "**" in stripped:
            parts = stripped.split("**")
            for j, part in enumerate(parts):
                if part:
                    run = p.add_run(part)
                    if j % 2 == 1:
                        run.bold = True
        else:
            p.add_run(stripped)
        i += 1; continue

    # Numbered items
    if line and len(line) > 2 and line[0].isdigit() and line[1] in ".)" or (len(line) > 3 and line[:2].isdigit() and line[2] == "."):
        p = doc.add_paragraph(style="List Number")
        text = line.lstrip("0123456789.) ").strip()
        if "**" in text:
            parts = text.split("**")
            for j, part in enumerate(parts):
                if part:
                    run = p.add_run(part)
                    if j % 2 == 1:
                        run.bold = True
        else:
            p.add_run(text)
        i += 1; continue

    # Empty line
    if not line.strip():
        i += 1; continue

    # Normal paragraph with bold handling
    p = doc.add_paragraph()
    text = line.strip()
    if "**" in text:
        parts = text.split("**")
        for j, part in enumerate(parts):
            if part:
                run = p.add_run(part)
                if j % 2 == 1:
                    run.bold = True
    else:
        p.add_run(text)
    i += 1

docx_path = SRC.with_suffix(".docx")
doc.save(str(docx_path))
print(f"Word: {docx_path}")

# ── HTML for PDF ──────────────────────────────────────────────────────
md_text = SRC.read_text(encoding="utf-8")
html_body = markdown.markdown(md_text, extensions=["tables", "fenced_code"])
html_full = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8">
<style>
  body {{ font-family: 'Segoe UI', Calibri, sans-serif; max-width: 900px; margin: auto; padding: 40px; line-height: 1.7; color: #222; }}
  h1 {{ color: #1b7a3d; border-bottom: 3px solid #1b7a3d; padding-bottom: 8px; }}
  h2 {{ color: #1b7a3d; margin-top: 2em; border-bottom: 1px solid #ddd; padding-bottom: 4px; }}
  h3 {{ color: #1a56db; }}
  blockquote {{ border-left: 4px solid #1b7a3d; margin: 1em 0; padding: 0.5em 1em; background: #f0faf4; color: #555; font-style: italic; }}
  table {{ border-collapse: collapse; width: 100%; margin: 1em 0; }}
  th, td {{ border: 1px solid #ccc; padding: 8px 12px; text-align: left; }}
  th {{ background: #1b7a3d; color: white; }}
  tr:nth-child(even) {{ background: #f8f8f8; }}
  code {{ background: #f0f0f0; padding: 2px 6px; border-radius: 3px; font-size: 0.95em; }}
  strong {{ color: #1a1a1a; }}
  ul, ol {{ margin: 0.5em 0; }}
  li {{ margin-bottom: 0.3em; }}
</style></head><body>{html_body}</body></html>"""

html_path = SRC.with_suffix(".html")
html_path.write_text(html_full, encoding="utf-8")
print(f"HTML: {html_path}")

# ── PDF via Edge ──────────────────────────────────────────────────────
pdf_path = SRC.with_suffix(".pdf")
edge = r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe"
if not os.path.exists(edge):
    edge = r"C:\Program Files\Microsoft\Edge\Application\msedge.exe"
subprocess.run([
    edge, "--headless", "--disable-gpu",
    f"--print-to-pdf={pdf_path}",
    "--no-pdf-header-footer",
    str(html_path.resolve())
], capture_output=True, timeout=30)
print(f"PDF:  {pdf_path}")
print("Done!")
