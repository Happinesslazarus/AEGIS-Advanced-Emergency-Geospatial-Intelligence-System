"""
Generate AEGIS-Complete-Overview.docx and AEGIS-Complete-Overview.pdf
from the markdown source using python-docx (Word) and HTML-to-PDF (built-in).
"""
import re, os, sys, html as html_mod
from pathlib import Path

SRC = Path(r"E:\aegis-v6-fullstack\docs\AEGIS-Complete-Overview.md")
OUT_DIR = SRC.parent
DOCX_OUT = OUT_DIR / "AEGIS-Complete-Overview.docx"
PDF_OUT  = OUT_DIR / "AEGIS-Complete-Overview.pdf"

md_text = SRC.read_text(encoding="utf-8")

# ───────────────────────── WORD DOCUMENT ─────────────────────────
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# -- styles
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)

style_h1 = doc.styles["Heading 1"]
style_h1.font.size = Pt(22)
style_h1.font.color.rgb = RGBColor(0x1A, 0x25, 0x6E)

style_h2 = doc.styles["Heading 2"]
style_h2.font.size = Pt(16)
style_h2.font.color.rgb = RGBColor(0x1A, 0x25, 0x6E)

style_h3 = doc.styles["Heading 3"]
style_h3.font.size = Pt(13)
style_h3.font.color.rgb = RGBColor(0x2D, 0x3A, 0x8C)

def add_md_table(lines):
    """Parse markdown table lines into a docx table."""
    rows = []
    for ln in lines:
        cells = [c.strip() for c in ln.strip().strip("|").split("|")]
        rows.append(cells)
    # skip separator row (row index 1)
    if len(rows) > 1:
        header = rows[0]
        data = [r for i, r in enumerate(rows[1:]) if not all(set(c) <= set("-: ") for c in r)]
    else:
        header = rows[0]
        data = []
    ncols = len(header)
    table = doc.add_table(rows=1 + len(data), cols=ncols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # header
    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    # data
    for ri, row in enumerate(data):
        for ci, val in enumerate(row):
            if ci < ncols:
                cell = table.rows[ri + 1].cells[ci]
                cell.text = val
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(10)

def inline_format(paragraph, text):
    """Handle **bold** and inline `code` in a paragraph."""
    parts = re.split(r'(\*\*.*?\*\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        else:
            paragraph.add_run(part)

lines = md_text.split("\n")
i = 0
in_code_block = False
code_block_lines = []

while i < len(lines):
    line = lines[i]

    # code blocks
    if line.strip().startswith("```"):
        if in_code_block:
            # end code block
            code_text = "\n".join(code_block_lines)
            p = doc.add_paragraph()
            run = p.add_run(code_text)
            run.font.name = "Consolas"
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
            pf = p.paragraph_format
            pf.space_before = Pt(4)
            pf.space_after = Pt(4)
            code_block_lines = []
            in_code_block = False
        else:
            in_code_block = True
            code_block_lines = []
        i += 1
        continue

    if in_code_block:
        code_block_lines.append(line)
        i += 1
        continue

    stripped = line.strip()

    # skip horizontal rules
    if stripped == "---":
        i += 1
        continue

    # headings
    if stripped.startswith("# ") and not stripped.startswith("## "):
        doc.add_heading(stripped[2:], level=0)
        i += 1
        continue
    if stripped.startswith("## "):
        doc.add_heading(stripped[3:], level=1)
        i += 1
        continue
    if stripped.startswith("### "):
        doc.add_heading(stripped[4:], level=2)
        i += 1
        continue
    if stripped.startswith("#### "):
        doc.add_heading(stripped[5:], level=3)
        i += 1
        continue

    # tables
    if "|" in stripped and stripped.startswith("|"):
        table_lines = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            table_lines.append(lines[i])
            i += 1
        add_md_table(table_lines)
        doc.add_paragraph()  # spacer
        continue

    # bullet points
    if stripped.startswith("- ") or stripped.startswith("* "):
        text = stripped[2:]
        p = doc.add_paragraph(style="List Bullet")
        inline_format(p, text)
        i += 1
        continue

    # numbered lists
    m = re.match(r'^(\d+)\.\s+(.*)', stripped)
    if m:
        text = m.group(2)
        p = doc.add_paragraph(style="List Number")
        inline_format(p, text)
        i += 1
        continue

    # empty line
    if not stripped:
        i += 1
        continue

    # normal paragraph
    p = doc.add_paragraph()
    inline_format(p, stripped)
    i += 1

doc.save(str(DOCX_OUT))
print(f"Word document saved: {DOCX_OUT}")

# ───────────────────────── PDF DOCUMENT ─────────────────────────
# Build HTML then use a lightweight approach
import markdown as md_lib

html_body = md_lib.markdown(md_text, extensions=["tables", "fenced_code"])

html_full = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>AEGIS — Complete System Overview</title>
<style>
@page {{ margin: 2cm; }}
body {{
    font-family: 'Segoe UI', Calibri, Arial, sans-serif;
    font-size: 11pt;
    line-height: 1.6;
    color: #1a1a2e;
    max-width: 800px;
    margin: 0 auto;
    padding: 20px;
}}
h1 {{ color: #1a256e; font-size: 24pt; border-bottom: 3px solid #1a256e; padding-bottom: 8px; }}
h2 {{ color: #1a256e; font-size: 18pt; border-bottom: 2px solid #ccc; padding-bottom: 4px; margin-top: 30px; }}
h3 {{ color: #2d3a8c; font-size: 14pt; margin-top: 20px; }}
h4 {{ color: #2d3a8c; font-size: 12pt; }}
table {{ border-collapse: collapse; width: 100%; margin: 10px 0; }}
th, td {{ border: 1px solid #ccc; padding: 6px 10px; text-align: left; font-size: 10pt; }}
th {{ background: #1a256e; color: white; font-weight: bold; }}
tr:nth-child(even) {{ background: #f4f6fb; }}
code {{ background: #f4f4f4; padding: 1px 4px; border-radius: 3px; font-family: Consolas, monospace; font-size: 9pt; color: #c0392b; }}
pre {{ background: #2c3e50; color: #ecf0f1; padding: 12px; border-radius: 6px; overflow-x: auto; font-size: 9pt; }}
pre code {{ background: none; color: #ecf0f1; }}
ul, ol {{ padding-left: 24px; }}
li {{ margin-bottom: 4px; }}
strong {{ color: #1a256e; }}
hr {{ border: none; border-top: 2px solid #1a256e; margin: 20px 0; }}
</style>
</head>
<body>
{html_body}
</body>
</html>
"""

html_path = OUT_DIR / "AEGIS-Complete-Overview.html"
html_path.write_text(html_full, encoding="utf-8")
print(f"HTML saved: {html_path}")

# Try to generate PDF
pdf_generated = False

# Method 1: Try weasyprint
try:
    from weasyprint import HTML as WeasyprintHTML
    WeasyprintHTML(string=html_full).write_pdf(str(PDF_OUT))
    pdf_generated = True
    print(f"PDF saved (weasyprint): {PDF_OUT}")
except ImportError:
    pass

# Method 2: Try pdfkit/wkhtmltopdf
if not pdf_generated:
    try:
        import pdfkit
        pdfkit.from_string(html_full, str(PDF_OUT))
        pdf_generated = True
        print(f"PDF saved (pdfkit): {PDF_OUT}")
    except Exception:
        pass

# Method 3: Use Edge/Chrome headless
if not pdf_generated:
    import subprocess
    for browser in [
        r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
        r"C:\Program Files\Microsoft\Edge\Application\msedge.exe",
        r"C:\Program Files\Google\Chrome\Application\chrome.exe",
        r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
    ]:
        if os.path.exists(browser):
            try:
                result = subprocess.run([
                    browser,
                    "--headless",
                    "--disable-gpu",
                    "--no-sandbox",
                    f"--print-to-pdf={PDF_OUT}",
                    str(html_path)
                ], capture_output=True, text=True, timeout=30)
                if PDF_OUT.exists() and PDF_OUT.stat().st_size > 1000:
                    pdf_generated = True
                    print(f"PDF saved (browser headless): {PDF_OUT}")
                    break
            except Exception as e:
                print(f"Browser PDF failed: {e}")
                continue

if not pdf_generated:
    print(f"\nPDF could not be generated automatically.")
    print(f"HTML file saved at: {html_path}")
    print(f"Open the HTML in a browser and use Print > Save as PDF.")

print("\nDone!")
