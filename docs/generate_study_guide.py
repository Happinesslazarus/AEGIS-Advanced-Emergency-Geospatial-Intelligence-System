"""Generate AEGIS-Study-Guide.docx and .pdf from markdown."""
import re, os
from pathlib import Path

SRC = Path(r"E:\aegis-v6-fullstack\docs\AEGIS-Study-Guide.md")
OUT_DIR = SRC.parent
DOCX_OUT = OUT_DIR / "AEGIS-Study-Guide.docx"
PDF_OUT  = OUT_DIR / "AEGIS-Study-Guide.pdf"

md_text = SRC.read_text(encoding="utf-8")

# ── WORD ──
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

for lvl, sname in [(1,"Heading 1"),(2,"Heading 2"),(3,"Heading 3")]:
    s = doc.styles[sname]
    s.font.color.rgb = RGBColor(0x0D,0x47,0xA1)
    s.font.size = Pt([0,22,16,13][lvl])

def add_md_table(lines):
    rows = []
    for ln in lines:
        cells = [c.strip() for c in ln.strip().strip("|").split("|")]
        rows.append(cells)
    header = rows[0]
    data = [r for i, r in enumerate(rows[1:]) if not all(set(c) <= set("-: ") for c in r)]
    ncols = len(header)
    t = doc.add_table(rows=1+len(data), cols=ncols)
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(header):
        cell = t.rows[0].cells[i]; cell.text = h
        for p in cell.paragraphs:
            for r in p.runs: r.bold = True; r.font.size = Pt(10)
    for ri, row in enumerate(data):
        for ci, val in enumerate(row):
            if ci < ncols:
                cell = t.rows[ri+1].cells[ci]; cell.text = val
                for p in cell.paragraphs:
                    for r in p.runs: r.font.size = Pt(10)

def inline_fmt(para, text):
    for part in re.split(r'(\*\*.*?\*\*|`[^`]+`)', text):
        if part.startswith("**") and part.endswith("**"):
            r = para.add_run(part[2:-2]); r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = para.add_run(part[1:-1]); r.font.name="Consolas"; r.font.size=Pt(9); r.font.color.rgb=RGBColor(0xC0,0x39,0x2B)
        else:
            para.add_run(part)

lines = md_text.split("\n")
i = 0; in_code = False; code_buf = []

while i < len(lines):
    line = lines[i]
    if line.strip().startswith("```"):
        if in_code:
            p = doc.add_paragraph()
            r = p.add_run("\n".join(code_buf))
            r.font.name="Consolas"; r.font.size=Pt(8); r.font.color.rgb=RGBColor(0x2C,0x3E,0x50)
            code_buf = []; in_code = False
        else:
            in_code = True; code_buf = []
        i += 1; continue
    if in_code:
        code_buf.append(line); i += 1; continue
    s = line.strip()
    if s == "---": i += 1; continue
    if s.startswith("# ") and not s.startswith("## "):
        doc.add_heading(s[2:], level=0); i += 1; continue
    if s.startswith("## "):
        doc.add_heading(s[3:], level=1); i += 1; continue
    if s.startswith("### "):
        doc.add_heading(s[4:], level=2); i += 1; continue
    if s.startswith("#### "):
        doc.add_heading(s[5:], level=3); i += 1; continue
    if "|" in s and s.startswith("|"):
        tl = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            tl.append(lines[i]); i += 1
        add_md_table(tl); doc.add_paragraph(); continue
    if s.startswith("- ") or s.startswith("* "):
        p = doc.add_paragraph(style="List Bullet"); inline_fmt(p, s[2:]); i += 1; continue
    m = re.match(r'^(\d+)\.\s+(.*)', s)
    if m:
        p = doc.add_paragraph(style="List Number"); inline_fmt(p, m.group(2)); i += 1; continue
    if not s: i += 1; continue
    p = doc.add_paragraph(); inline_fmt(p, s); i += 1

doc.save(str(DOCX_OUT))
print(f"Word: {DOCX_OUT}")

# ── PDF via HTML + Edge headless ──
import markdown as md_lib
html_body = md_lib.markdown(md_text, extensions=["tables","fenced_code"])
html_full = f"""<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8">
<title>AEGIS Study Guide</title><style>
@page {{ margin: 2cm; }}
body {{ font-family: 'Segoe UI', Calibri, sans-serif; font-size: 11pt; line-height: 1.7; color: #1a1a2e; max-width: 820px; margin: 0 auto; padding: 20px; }}
h1 {{ color: #0d47a1; font-size: 26pt; border-bottom: 3px solid #0d47a1; padding-bottom: 8px; }}
h2 {{ color: #0d47a1; font-size: 19pt; border-bottom: 2px solid #ddd; padding-bottom: 4px; margin-top: 32px; }}
h3 {{ color: #1565c0; font-size: 14pt; margin-top: 22px; }}
h4 {{ color: #1565c0; font-size: 12pt; }}
table {{ border-collapse: collapse; width: 100%; margin: 12px 0; }}
th, td {{ border: 1px solid #ccc; padding: 7px 12px; text-align: left; font-size: 10pt; }}
th {{ background: #0d47a1; color: white; }}
tr:nth-child(even) {{ background: #f0f4ff; }}
code {{ background: #f4f4f4; padding: 2px 5px; border-radius: 3px; font-family: Consolas, monospace; font-size: 9.5pt; color: #c0392b; }}
pre {{ background: #263238; color: #eceff1; padding: 14px; border-radius: 8px; overflow-x: auto; font-size: 9pt; line-height: 1.5; }}
pre code {{ background: none; color: #eceff1; }}
ul, ol {{ padding-left: 24px; }}
li {{ margin-bottom: 5px; }}
strong {{ color: #0d47a1; }}
hr {{ border: none; border-top: 2px solid #0d47a1; margin: 24px 0; }}
blockquote {{ border-left: 4px solid #0d47a1; margin: 12px 0; padding: 8px 16px; background: #f0f4ff; }}
</style></head><body>{html_body}</body></html>"""

html_path = OUT_DIR / "AEGIS-Study-Guide.html"
html_path.write_text(html_full, encoding="utf-8")

import subprocess
pdf_done = False
for browser in [
    r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
    r"C:\Program Files\Microsoft\Edge\Application\msedge.exe",
    r"C:\Program Files\Google\Chrome\Application\chrome.exe",
]:
    if os.path.exists(browser):
        try:
            subprocess.run([browser,"--headless","--disable-gpu","--no-sandbox",
                f"--print-to-pdf={PDF_OUT}", str(html_path)],
                capture_output=True, text=True, timeout=30)
            if PDF_OUT.exists() and PDF_OUT.stat().st_size > 1000:
                pdf_done = True; print(f"PDF: {PDF_OUT}"); break
        except: continue

if not pdf_done:
    print(f"PDF failed. Open {html_path} in browser → Print → Save as PDF")

print("Done!")
